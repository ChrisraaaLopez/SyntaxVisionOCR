"""
syntax_vision.py  ·  v4.4
==========================
OCR → PKM Markdown | Dual-engine: Tesseract + EasyOCR | IA Multi-proveedor

Cambios en v4.4:
  - BUGFIX CRÍTICO: mask_graphic_regions() ahora está DEFINIDA (en v4.3 era
    llamada pero nunca definida, causando NameError silencioso y saltando
    el enmascarado de gráficos → texto perdido en sellos/logos).
  - IA MULTIMODAL: los proveedores que soportan visión (Claude, Gemini,
    GPT-4o) ahora reciben IMAGEN + TEXTO OCR juntos. El modelo compara
    ambos y reconstruye el texto faltante o mal reconocido, dando resultados
    mucho más precisos en documentos con sellos, manchas o fondos complejos.
  - DeepSeek mantiene modo texto-only (su API no soporta visión aún).
  - Modos disponibles: auto, photo, document, raw.

Dependencias:
    pip install PyQt6 pytesseract opencv-python Pillow python-docx reportlab easyocr

Tesseract:
    https://github.com/UB-Mannheim/tesseract/wiki  (incluir langpack spa)
Montserrat: descarga las 4 variantes .ttf desde Google Fonts y colócalas en:
    Windows : C:/Windows/Fonts/
    Linux   : ~/.local/share/fonts/
    macOS   : ~/Library/Fonts/
"""

import sys
import os
import re
import io
import json
import base64
import sqlite3
import urllib.request
import cv2
import pytesseract
import numpy as np
from PIL import Image, ImageEnhance, ImageFilter
from datetime import datetime
from collections import deque
from PyQt6.QtWidgets import QApplication
from PyQt6.QtGui import QIcon

_easyocr_reader = None
_easyocr_langs_loaded = None

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable
)

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QSplitter,
    QPushButton, QVBoxLayout, QHBoxLayout, QLabel,
    QTextEdit, QListWidget, QListWidgetItem, QFileDialog,
    QMessageBox, QFrame, QInputDialog, QDialog,
    QComboBox, QGroupBox, QLineEdit, QSizePolicy
)
from PyQt6.QtGui import QPixmap, QImage, QFontDatabase, QIcon, QTextCursor
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize


# ══════════════════════════════════════════════════════════════════
#  RUTAS Y CONSTANTES
# ══════════════════════════════════════════════════════════════════

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH   = os.path.join(BASE_DIR, "resources", "logo.png")
DB_PATH     = os.path.join(BASE_DIR, "syntax_vision_history.db")
MAX_CHARS   = 10_000
MAX_HISTORY = 20

def resource_path(relative_path):
    """ Obtiene la ruta absoluta de los recursos para el EXE """
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

tesseract_folder = resource_path("Tesseract-OCR")
pytesseract.pytesseract.tesseract_cmd = os.path.join(tesseract_folder, "tesseract.exe")
os.environ["EASYOCR_MODULE_PATH"] = resource_path(os.path.join("resources", "models"))


# ══════════════════════════════════════════════════════════════════
#  CACHÉ DE IMAGEN PREPROCESADA
# ══════════════════════════════════════════════════════════════════

_preproc_cache: dict = {"path": None, "mode": None, "image": None}

def get_preprocessed(pil_img: Image.Image, path: str, mode: str) -> Image.Image:
    """Devuelve imagen preprocesada desde caché si path+mode coinciden."""
    if _preproc_cache["path"] == path and _preproc_cache["mode"] == mode:
        return _preproc_cache["image"]
    processed = preprocess_for_ocr(pil_img, mode=mode)
    _preproc_cache["path"]  = path
    _preproc_cache["mode"]  = mode
    _preproc_cache["image"] = processed
    return processed


# ══════════════════════════════════════════════════════════════════
#  UTILIDAD: IMAGEN → BASE64  (para IA multimodal)
# ══════════════════════════════════════════════════════════════════

def pil_to_base64(pil_img: Image.Image, max_side: int = 1600) -> tuple[str, str]:
    """
    Convierte PIL Image a base64 JPEG para enviar a APIs de visión.
    Escala la imagen si algún lado supera max_side para reducir tokens/costo.
    Retorna (base64_string, media_type).
    """
    img = pil_img.convert("RGB")
    w, h = img.size
    if max(w, h) > max_side:
        scale = max_side / max(w, h)
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return b64, "image/jpeg"


# ══════════════════════════════════════════════════════════════════
#  TEMAS DE COLOR
# ══════════════════════════════════════════════════════════════════

THEMES = {
    "Industrial Dark": {
        "bg_root":   "#06141B",
        "bg_panel":  "#11212D",
        "bg_widget": "#253745",
        "accent":    "#4A5C6A",
        "border":    "#4A5C6A",
        "text_main": "#CCD0CF",
        "text_sub":  "#9BA8AB",
        "btn_hover": "#4A5C6A",
        "highlight": "#9BA8AB",
    },
    "Nordic Night": {
        "bg_root":   "#2E3440",
        "bg_panel":  "#3B4252",
        "bg_widget": "#434C5E",
        "accent":    "#88C0D0",
        "border":    "#4C566A",
        "text_main": "#ECEFF4",
        "text_sub":  "#D8DEE9",
        "btn_hover": "#5E81AC",
        "highlight": "#88C0D0",
    },
    "Solarized": {
        "bg_root":   "#002B36",
        "bg_panel":  "#073642",
        "bg_widget": "#094652",
        "accent":    "#2AA198",
        "border":    "#586E75",
        "text_main": "#FDF6E3",
        "text_sub":  "#EEE8D5",
        "btn_hover": "#268BD2",
        "highlight": "#2AA198",
    },
    "Monokai": {
        "bg_root":   "#1C1C1C",
        "bg_panel":  "#272822",
        "bg_widget": "#3E3D32",
        "accent":    "#F92672",
        "border":    "#49483E",
        "text_main": "#F8F8F2",
        "text_sub":  "#CFCFC2",
        "btn_hover": "#A6E22E",
        "highlight": "#E6DB74",
    },
    "Cyberpunk": {
        "bg_root":   "#0D0D0D",
        "bg_panel":  "#1A0A2E",
        "bg_widget": "#16213E",
        "accent":    "#00FFF0",
        "border":    "#7B2FBE",
        "text_main": "#E0E0E0",
        "text_sub":  "#A0A0C0",
        "btn_hover": "#7B2FBE",
        "highlight": "#00FFF0",
    },
    "Papel Blanco": {
        "bg_root":   "#F5F5F0",
        "bg_panel":  "#FFFFFF",
        "bg_widget": "#FAFAFA",
        "accent":    "#3B82F6",
        "border":    "#D1D5DB",
        "text_main": "#111827",
        "text_sub":  "#6B7280",
        "btn_hover": "#2563EB",
        "highlight": "#3B82F6",
    },
}

CURRENT_THEME = "Industrial Dark"


def build_stylesheet(theme_name: str) -> str:
    t = THEMES[theme_name]
    return f"""
    QMainWindow, QDialog {{
        background-color: {t['bg_root']};
    }}
    QWidget {{
        background-color: {t['bg_root']};
        color: {t['text_main']};
        font-family: 'Montserrat', 'Segoe UI', sans-serif;
        font-size: 13px;
    }}
    QSplitter::handle {{
        background-color: {t['border']};
        width: 2px;
        height: 2px;
    }}
    QPushButton {{
        background-color: {t['bg_widget']};
        color: {t['text_main']};
        border: 1px solid {t['border']};
        border-radius: 6px;
        padding: 8px 14px;
        font-weight: 600;
        font-family: 'Montserrat', 'Segoe UI', sans-serif;
    }}
    QPushButton:hover {{
        background-color: {t['btn_hover']};
        color: #ffffff;
        border-color: {t['highlight']};
    }}
    QPushButton:pressed {{
        background-color: {t['accent']};
    }}
    QPushButton:disabled {{
        background-color: {t['bg_panel']};
        color: {t['text_sub']};
        border: 1px solid {t['bg_widget']};
    }}
    QPushButton#AccentButton {{
        background-color: {t['accent']};
        border: 1px solid {t['highlight']};
        color: #ffffff;
    }}
    QPushButton#AccentButton:hover {{
        background-color: {t['highlight']};
        color: {t['bg_root']};
    }}
    QPushButton#DangerButton {{
        background-color: transparent;
        border: 1px solid #7a3a3a;
        color: #f08080;
    }}
    QPushButton#DangerButton:hover {{
        background-color: #7a3a3a;
        color: #ffffff;
    }}
    QLabel {{
        color: {t['text_sub']};
        font-size: 11px;
        font-weight: 700;
        letter-spacing: 1px;
        background: transparent;
        font-family: 'Montserrat', 'Segoe UI', sans-serif;
    }}
    QLabel#SectionTitle {{
        color: {t['text_main']};
        font-size: 11px;
        padding: 4px 2px;
        border-bottom: 1px solid {t['border']};
        letter-spacing: 2px;
    }}
    QLabel#AppTitle {{
        color: {t['highlight']};
        font-size: 18px;
        font-weight: 800;
        letter-spacing: 3px;
        background: transparent;
    }}
    QLabel#AppSubtitle {{
        color: {t['text_sub']};
        font-size: 10px;
        font-weight: 400;
        letter-spacing: 2px;
        background: transparent;
    }}
    QLabel#CharCounter {{
        font-size: 11px;
        font-weight: 700;
        font-family: 'Montserrat', 'Consolas', monospace;
        padding: 2px 8px;
        border-radius: 4px;
    }}
    QTextEdit {{
        background-color: {t['bg_panel']};
        color: {t['text_main']};
        border: 1px solid {t['border']};
        border-radius: 6px;
        font-family: 'Montserrat', 'Consolas', monospace;
        font-size: 13px;
        padding: 8px;
        selection-background-color: {t['accent']};
    }}
    QTextEdit:focus {{
        border: 1px solid {t['highlight']};
    }}
    QListWidget {{
        background-color: {t['bg_panel']};
        border: 1px solid {t['border']};
        border-radius: 6px;
        color: {t['text_main']};
        font-family: 'Montserrat', 'Segoe UI', sans-serif;
    }}
    QListWidget::item {{
        padding: 7px 10px;
        border-bottom: 1px solid {t['bg_widget']};
    }}
    QListWidget::item:selected {{
        background-color: {t['accent']};
        color: #ffffff;
        border-radius: 4px;
    }}
    QListWidget::item:hover {{
        background-color: {t['bg_widget']};
    }}
    QLabel#DropArea {{
        border: 2px dashed {t['border']};
        background-color: {t['bg_panel']};
        border-radius: 10px;
        color: {t['text_sub']};
        font-size: 12px;
        font-weight: 500;
        min-height: 170px;
    }}
    QLabel#CameraView {{
        background-color: {t['bg_root']};
        border: 2px solid {t['border']};
        border-radius: 8px;
    }}
    QComboBox {{
        background-color: {t['bg_widget']};
        color: {t['text_main']};
        border: 1px solid {t['border']};
        border-radius: 5px;
        padding: 5px 10px;
        font-family: 'Montserrat', 'Segoe UI', sans-serif;
        font-weight: 600;
    }}
    QComboBox QAbstractItemView {{
        background-color: {t['bg_panel']};
        color: {t['text_main']};
        selection-background-color: {t['accent']};
        border: 1px solid {t['border']};
    }}
    QGroupBox {{
        border: 1px solid {t['border']};
        border-radius: 6px;
        margin-top: 8px;
        font-family: 'Montserrat', 'Segoe UI', sans-serif;
        font-size: 10px;
        font-weight: 700;
        color: {t['text_sub']};
        letter-spacing: 1px;
    }}
    QGroupBox::title {{
        subcontrol-origin: margin;
        left: 10px;
        padding: 0 4px;
    }}
    QScrollBar:vertical {{
        background: {t['bg_root']};
        width: 7px;
        border-radius: 3px;
    }}
    QScrollBar::handle:vertical {{
        background: {t['accent']};
        border-radius: 3px;
        min-height: 20px;
    }}
    QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0; }}
    QScrollBar:horizontal {{
        background: {t['bg_root']};
        height: 7px;
    }}
    QScrollBar::handle:horizontal {{
        background: {t['accent']};
        border-radius: 3px;
    }}
    QFrame[frameShape="4"], QFrame[frameShape="5"] {{
        color: {t['border']};
    }}
    """


# ══════════════════════════════════════════════════════════════════
#  FUENTE MONTSERRAT
# ══════════════════════════════════════════════════════════════════

def register_montserrat():
    search_dirs = [
        r"C:\Windows\Fonts",
        os.path.expanduser(r"~\AppData\Local\Microsoft\Windows\Fonts"),
        os.path.expanduser("~/.local/share/fonts"),
        "/usr/share/fonts/truetype",
        "/usr/local/share/fonts",
        os.path.expanduser("~/Library/Fonts"),
        "/Library/Fonts",
        "/System/Library/Fonts",
    ]
    found = 0
    for folder in search_dirs:
        if not os.path.isdir(folder):
            continue
        for fname in os.listdir(folder):
            if "montserrat" in fname.lower() and fname.endswith(".ttf"):
                QFontDatabase.addApplicationFont(os.path.join(folder, fname))
                found += 1
    if found == 0:
        print("[INFO] Montserrat no encontrada; usando Segoe UI / sans-serif.")
    else:
        print(f"[INFO] Montserrat registrada ({found} variantes).")


# ══════════════════════════════════════════════════════════════════
#  BASE DE DATOS SQLITE — HISTORIAL PERSISTENTE
# ══════════════════════════════════════════════════════════════════

def db_init() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS history (
            id        INTEGER PRIMARY KEY AUTOINCREMENT,
            title     TEXT    NOT NULL,
            content   TEXT    NOT NULL,
            timestamp TEXT    NOT NULL
        )
    """)
    conn.commit()
    return conn


def db_save_entry(conn: sqlite3.Connection, title: str, content: str):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute(
        "INSERT INTO history (title, content, timestamp) VALUES (?, ?, ?)",
        (title, content, ts)
    )
    conn.execute(f"""
        DELETE FROM history
        WHERE id NOT IN (
            SELECT id FROM history
            ORDER BY id DESC
            LIMIT {MAX_HISTORY}
        )
    """)
    conn.commit()


def db_load_all(conn: sqlite3.Connection) -> list[dict]:
    cur = conn.execute(
        "SELECT id, title, content, timestamp FROM history ORDER BY id DESC"
    )
    return [
        {"id": r[0], "title": r[1], "content": r[2], "timestamp": r[3]}
        for r in cur.fetchall()
    ]


def db_clear(conn: sqlite3.Connection):
    conn.execute("DELETE FROM history")
    conn.commit()


# ══════════════════════════════════════════════════════════════════
#  DETECCIÓN DE CÁMARAS
# ══════════════════════════════════════════════════════════════════

class CameraDetectThread(QThread):
    cameras_found = pyqtSignal(list)

    def run(self):
        available = []
        backend = cv2.CAP_DSHOW if os.name == "nt" else cv2.CAP_ANY
        for i in range(6):
            try:
                cap = cv2.VideoCapture(i, backend)
                if not cap.isOpened():
                    cap.release()
                    continue
                ret, _ = cap.read()
                if ret:
                    w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                    h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                    if i == 0:
                        label = f"Cámara integrada  [{w}×{h}]"
                    else:
                        label = f"Cámara USB {i}  [{w}×{h}]"
                    available.append({"index": i, "label": label})
                cap.release()
            except Exception:
                pass
        if not available:
            available = [{"index": 0, "label": "Cámara 0 (predeterminada)"}]
        self.cameras_found.emit(available)


# ══════════════════════════════════════════════════════════════════
#  REORDENAMIENTO DE TEXTO POR BOUNDING BOXES
# ══════════════════════════════════════════════════════════════════

_KEEP_SHORT_TOKENS = {
    "la", "el", "en", "de", "se", "le", "al", "lo", "un", "su",
    "las", "los", "una", "del", "por", "con", "que",
    "La", "El", "En", "De", "Se", "Le", "Al", "Lo", "Un", "Su",
}


def _extract_zone_conf(img_arr: np.ndarray, conf_threshold: int = 40,
                       psm: int = 6, lang: str = "spa") -> str:
    img_h = img_arr.shape[0]
    LINE_TOL = max(10, int(img_h * 0.018))
    try:
        pil_zone = Image.fromarray(img_arr)
        data = pytesseract.image_to_data(
            pil_zone, lang=lang,
            config=f"--oem 3 --psm {psm}",
            output_type=pytesseract.Output.DICT
        )
    except Exception:
        return ""

    n = len(data["text"])
    words = []
    for i in range(n):
        word = data["text"][i].strip()
        conf_raw = data["conf"][i]
        conf = int(conf_raw) if str(conf_raw).lstrip("-").isdigit() else -1
        if not word or conf < conf_threshold:
            continue
        words.append({
            "text":  word,
            "left":  int(data["left"][i]),
            "top":   int(data["top"][i]),
            "block": int(data["block_num"][i]),
            "par":   int(data["par_num"][i]),
            "line":  int(data["line_num"][i]),
        })
    if not words:
        return ""

    line_groups: dict = {}
    for w in words:
        key = (w["block"], w["par"], w["line"])
        line_groups.setdefault(key, []).append(w)

    line_meta: dict = {}
    for key, wlist in line_groups.items():
        line_meta[key] = {
            "top":   min(x["top"]  for x in wlist),
            "left":  min(x["left"] for x in wlist),
            "words": sorted(wlist, key=lambda x: x["left"]),
            "block": wlist[0]["block"],
        }

    sorted_keys = sorted(line_meta.keys(),
                         key=lambda k: (line_meta[k]["top"], line_meta[k]["left"]))
    rows: list = []
    cur: list = []
    last_top = None
    for key in sorted_keys:
        top = line_meta[key]["top"]
        if last_top is None or abs(top - last_top) <= LINE_TOL:
            cur.append(key)
            last_top = min(last_top, top) if last_top is not None else top
        else:
            if cur:
                rows.append(cur)
            cur = [key]
            last_top = top
    if cur:
        rows.append(cur)

    out_lines: list = []
    prev_block = None
    for row in rows:
        for key in sorted(row, key=lambda k: line_meta[k]["left"]):
            meta = line_meta[key]
            if prev_block is not None and meta["block"] != prev_block:
                out_lines.append("")
            out_lines.append(" ".join(x["text"] for x in meta["words"]))
            prev_block = meta["block"]
    return "\n".join(out_lines)


def _post_filter_ocr_lines(text: str) -> str:
    lines = text.splitlines()
    out = []
    for line in lines:
        s = line.strip()
        if not s:
            out.append('')
            continue
        tokens = s.split()
        clean_tokens = [t.strip("'\".,;:!?-_ ") for t in tokens]

        if (all(len(t) <= 2 for t in clean_tokens)
                and not any(t in _KEEP_SHORT_TOKENS for t in tokens)):
            continue

        while tokens:
            t = tokens[0]
            t_clean = t.strip("'\".,;:!?-_ ")
            if not t_clean:
                tokens = tokens[1:]
                continue
            if (len(t_clean) == 1 and t_clean.islower()
                    and t not in _KEEP_SHORT_TOKENS):
                tokens = tokens[1:]
                continue
            if not any(c.isalpha() for c in t_clean):
                tokens = tokens[1:]
                continue
            break

        if not tokens:
            continue

        tokens[0] = re.sub(r'^\d+([A-ZÁÉÍÓÚÑa-záéíóúñ])', r'\1', tokens[0])
        if not tokens[0]:
            tokens = tokens[1:]
        if not tokens:
            continue

        out.append(" ".join(tokens))

    result: list = []
    blanks = 0
    for l in out:
        if l == '':
            blanks += 1
            if blanks <= 1:
                result.append(l)
        else:
            blanks = 0
            result.append(l)
    return '\n'.join(result).strip()


def reorder_ocr_text_by_layout(pil_img: Image.Image, lang: str = "spa") -> str:
    try:
        data = pytesseract.image_to_data(
            pil_img, lang=lang, config="--oem 3 --psm 3",
            output_type=pytesseract.Output.DICT
        )
    except Exception:
        return ""

    n = len(data["text"])
    if n == 0:
        return ""

    raw_confs = []
    for i in range(n):
        c = data["conf"][i]
        val = int(c) if str(c).lstrip("-").isdigit() else -1
        if val > 0:
            raw_confs.append(val)
    adaptive_threshold = max(45, int(np.percentile(raw_confs, 25))) if raw_confs else 45

    words = []
    char_heights = []
    for i in range(n):
        word = data["text"][i].strip()
        conf_raw = data["conf"][i]
        conf = int(conf_raw) if str(conf_raw).lstrip("-").isdigit() else -1
        if not word or conf < adaptive_threshold:
            continue
        if len(word) == 1 and not word.isalnum():
            continue
        alnum_ratio = sum(1 for c in word if c.isalnum()) / max(len(word), 1)
        if alnum_ratio < 0.35 and len(word) > 2:
            continue
        h_char = int(data["height"][i])
        if h_char > 0:
            char_heights.append(h_char)
        words.append({
            "text":  word,
            "left":  int(data["left"][i]),
            "top":   int(data["top"][i]),
            "block": int(data["block_num"][i]),
            "par":   int(data["par_num"][i]),
            "line":  int(data["line_num"][i]),
        })

    if not words:
        return ""

    if char_heights:
        median_h = float(np.median(char_heights))
        LINE_TOLERANCE = max(10, int(median_h * 1.2))
    else:
        LINE_TOLERANCE = 18

    line_groups: dict[tuple, list] = {}
    for w in words:
        key = (w["block"], w["par"], w["line"])
        line_groups.setdefault(key, []).append(w)

    line_meta = {}
    for key, wlist in line_groups.items():
        line_meta[key] = {
            "top":   min(w["top"]  for w in wlist),
            "left":  min(w["left"] for w in wlist),
            "words": sorted(wlist, key=lambda w: w["left"]),
            "block": wlist[0]["block"],
        }

    sorted_keys = sorted(line_meta.keys(),
                         key=lambda k: (line_meta[k]["top"], line_meta[k]["left"]))
    row_groups: list[list] = []
    current_row: list = []
    last_top = None

    for key in sorted_keys:
        top = line_meta[key]["top"]
        if last_top is None or abs(top - last_top) <= LINE_TOLERANCE:
            current_row.append(key)
            last_top = min(last_top, top) if last_top is not None else top
        else:
            if current_row:
                row_groups.append(current_row)
            current_row = [key]
            last_top = top
    if current_row:
        row_groups.append(current_row)

    output_lines = []
    prev_block = None
    for row in row_groups:
        row_sorted = sorted(row, key=lambda k: line_meta[k]["left"])
        for key in row_sorted:
            meta = line_meta[key]
            block_num = meta["block"]
            if prev_block is not None and block_num != prev_block:
                output_lines.append("")
            output_lines.append(" ".join(w["text"] for w in meta["words"]))
            prev_block = block_num

    return "\n".join(output_lines)


# ══════════════════════════════════════════════════════════════════
#  PREPROCESAMIENTO DE IMAGEN
# ══════════════════════════════════════════════════════════════════

def mask_graphic_regions(gray: np.ndarray) -> np.ndarray:
    """
    v4.4 — FUNCIÓN PREVIAMENTE FALTANTE (BUG en v4.3).
    Detecta y blanquea zonas con alta densidad de bordes Sobel
    (logos, firmas, sellos húmedos) ANTES del umbral adaptativo.
    Esto evita que Tesseract intente leer gráficos como texto.
    """
    sobelx = cv2.Sobel(gray, cv2.CV_64F, 1, 0, ksize=3)
    sobely = cv2.Sobel(gray, cv2.CV_64F, 0, 1, ksize=3)
    magnitude = np.sqrt(sobelx**2 + sobely**2)
    magnitude = np.uint8(np.clip(magnitude / magnitude.max() * 255, 0, 255))

    # Dividir en bloques y calcular densidad de bordes por bloque
    h, w = gray.shape
    block_h = max(1, h // 20)
    block_w = max(1, w // 20)
    result = gray.copy()

    for y in range(0, h, block_h):
        for x in range(0, w, block_w):
            block = magnitude[y:y+block_h, x:x+block_w]
            # Si más del 35% de píxeles son bordes fuertes → probable gráfico
            edge_density = np.sum(block > 80) / max(block.size, 1)
            if edge_density > 0.35:
                # Verificar que no sea texto denso (texto tiene bordes pero
                # también mucho blanco alrededor)
                gray_block = gray[y:y+block_h, x:x+block_w]
                white_ratio = np.sum(gray_block > 200) / max(gray_block.size, 1)
                # Si hay pocos píxeles blancos Y alta densidad de borde → gráfico
                if white_ratio < 0.30:
                    result[y:y+block_h, x:x+block_w] = 255
    return result


def mask_non_text_regions(gray: np.ndarray) -> np.ndarray:
    h, w = gray.shape
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    num_labels, labels, stats, _ = cv2.connectedComponentsWithStats(binary, connectivity=8)
    result = gray.copy()
    min_char_area = 30
    graphic_area_threshold = (h * w) * 0.005

    for i in range(1, num_labels):
        x, y, bw, bh, area = (stats[i, cv2.CC_STAT_LEFT],
                               stats[i, cv2.CC_STAT_TOP],
                               stats[i, cv2.CC_STAT_WIDTH],
                               stats[i, cv2.CC_STAT_HEIGHT],
                               stats[i, cv2.CC_STAT_AREA])
        if area < min_char_area:
            continue
        if area < graphic_area_threshold:
            continue
        aspect  = bw / max(bh, 1)
        density = area / max(bw * bh, 1)
        is_graphic = False
        if area > (h * w) * 0.008 and density < 0.25:
            is_graphic = True
        if area > (h * w) * 0.008 and 0.6 < aspect < 1.7 and density > 0.35:
            is_graphic = True
        if aspect > 5.0 and bh < h * 0.04 and area > (h * w) * 0.003:
            is_graphic = True
        if is_graphic:
            pad = 4
            x1 = max(0, x - pad); y1 = max(0, y - pad)
            x2 = min(w, x + bw + pad); y2 = min(h, y + bh + pad)
            result[y1:y2, x1:x2] = 255
    return result


def _correct_perspective(gray: np.ndarray) -> np.ndarray:
    h, w = gray.shape
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(blurred, 50, 150)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
    edges = cv2.dilate(edges, kernel, iterations=1)
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        return gray
    contours = sorted(contours, key=cv2.contourArea, reverse=True)
    doc_contour = None
    for cnt in contours[:5]:
        peri = cv2.arcLength(cnt, True)
        approx = cv2.approxPolyDP(cnt, 0.02 * peri, True)
        if len(approx) == 4 and cv2.contourArea(approx) > 0.20 * h * w:
            doc_contour = approx
            break
    if doc_contour is None:
        return gray
    pts = doc_contour.reshape(4, 2).astype(np.float32)
    rect = np.zeros((4, 2), dtype=np.float32)
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]
    widthA  = np.linalg.norm(rect[2] - rect[3])
    widthB  = np.linalg.norm(rect[1] - rect[0])
    heightA = np.linalg.norm(rect[1] - rect[2])
    heightB = np.linalg.norm(rect[0] - rect[3])
    dst_w = int(max(widthA, widthB))
    dst_h = int(max(heightA, heightB))
    if dst_w < 100 or dst_h < 100:
        return gray
    dst = np.array([[0,0],[dst_w-1,0],[dst_w-1,dst_h-1],[0,dst_h-1]], dtype=np.float32)
    M = cv2.getPerspectiveTransform(rect, dst)
    return cv2.warpPerspective(gray, M, (dst_w, dst_h))


def preprocess_for_ocr(pil_img: Image.Image, mode: str = "auto") -> Image.Image:
    """
    Modos disponibles: auto, photo, document, raw.
    """
    if mode == "raw":
        return pil_img.convert("L") if pil_img.mode != "L" else pil_img
    if pil_img.mode not in ("RGB", "L"):
        pil_img = pil_img.convert("RGB")
    img_np = np.array(pil_img.convert("RGB"))
    gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)

    if mode == "auto":
        mode = "photo"

    if mode == "document":
        gray = _correct_perspective(gray)
        h, w = gray.shape
        if w < 2800:
            scale = max(2, 2800 // w)
            gray = cv2.resize(gray, (w * scale, h * scale), interpolation=cv2.INTER_CUBIC)
        clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(8, 8))
        equalized = clahe.apply(gray)
        denoised = cv2.fastNlMeansDenoising(equalized, h=10, templateWindowSize=7,
                                             searchWindowSize=21)
        blur = cv2.GaussianBlur(denoised, (0, 0), 1.5)
        sharp = cv2.addWeighted(denoised, 1.6, blur, -0.6, 0)
        # mask_graphic_regions ahora está definida correctamente (bugfix v4.4)
        masked_graphic = mask_graphic_regions(sharp)
        masked = mask_non_text_regions(masked_graphic)
        binary = cv2.adaptiveThreshold(
            masked, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, blockSize=51, C=15
        )
        k_close = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
        binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, k_close)
        return Image.fromarray(binary)

    else:  # photo
        gray = _correct_perspective(gray)
        h, w = gray.shape
        if w < 1800:
            scale = max(2, 1800 // w)
            gray = cv2.resize(gray, (w * scale, h * scale), interpolation=cv2.INTER_CUBIC)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        equalized = clahe.apply(gray)
        denoised = cv2.bilateralFilter(equalized, 9, 75, 75)
        blur = cv2.GaussianBlur(denoised, (0, 0), 2)
        sharp = cv2.addWeighted(denoised, 1.8, blur, -0.8, 0)
        binary = cv2.adaptiveThreshold(
            sharp, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, blockSize=31, C=10
        )
        k_open = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
        binary = cv2.morphologyEx(binary, cv2.MORPH_OPEN, k_open)
        return Image.fromarray(binary)


def run_ocr_multilang(pil_img: Image.Image, lang: str = "spa", psm: int = 6) -> str:
    config = f"--oem 3 --psm {psm}"
    try:
        text = pytesseract.image_to_string(pil_img, lang=lang, config=config)
    except pytesseract.TesseractError:
        text = pytesseract.image_to_string(pil_img, lang=lang)
    if len(text.strip()) < 20:
        try:
            t2 = pytesseract.image_to_string(pil_img, lang=lang, config="--oem 3 --psm 11")
            if len(t2.strip()) > len(text.strip()):
                text = t2
        except Exception:
            pass
    return text


def run_ocr_easyocr(pil_img: Image.Image, lang: str = "spa") -> str:
    global _easyocr_reader, _easyocr_langs_loaded
    lang_map = {
        "spa": ["es"], "eng": ["en"], "spa+eng": ["es", "en"],
        "fra": ["fr"], "deu": ["de"], "ita": ["it"], "por": ["pt"],
    }
    easy_langs = lang_map.get(lang, ["es"])
    if _easyocr_reader is None or _easyocr_langs_loaded != easy_langs:
        try:
            import easyocr
            _easyocr_reader = easyocr.Reader(easy_langs, gpu=False, verbose=False)
            _easyocr_langs_loaded = easy_langs
        except ImportError:
            raise RuntimeError("EasyOCR no está instalado.\nEjecuta:  pip install easyocr")
    img_np = np.array(pil_img.convert("RGB"))
    results = _easyocr_reader.readtext(img_np, detail=0, paragraph=True)
    return "\n\n".join(str(r) for r in results if str(r).strip())


def run_ocr(pil_img: Image.Image, engine: str = "tesseract",
            lang: str = "spa", mode: str = "auto") -> str:
    if engine == "easyocr":
        return run_ocr_easyocr(pil_img, lang=lang)
    psm = 3 if mode == "document" else 6
    return run_ocr_multilang(pil_img, lang=lang, psm=psm)


# ══════════════════════════════════════════════════════════════════
#  LIMPIEZA DE TEXTO OCR
# ══════════════════════════════════════════════════════════════════

def clean_ocr_text(raw: str) -> str:
    raw = raw.replace('\u2014', '-').replace('\u2013', '-')
    raw = raw.replace('\u201c', '"').replace('\u201d', '"')
    raw = raw.replace('\u2018', "'").replace('\u2019', "'")

    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\xA0-\xFF]', '', raw)
    text = re.sub(r'[\xa6\xa7\x5e~`|\\]', '', text)
    lines_raw = text.splitlines()
    lines_clean = []
    for line in lines_raw:
        stripped = line.strip()
        if not stripped:
            lines_clean.append('')
            continue
        alnum_count = sum(1 for c in stripped if c.isalnum())
        alnum_ratio = alnum_count / max(len(stripped), 1)
        if alnum_ratio < 0.40 and len(stripped) > 3:
            continue
        if len(stripped) <= 2 and alnum_count == 0:
            continue
        inner = re.sub(r'(?<=\S) {2,}', ' ', line.rstrip())
        lines_clean.append(inner)

    result = []
    prev_blank = False
    for line in lines_clean:
        if line.strip() == '':
            if not prev_blank:
                result.append(line)
            prev_blank = True
        else:
            prev_blank = False
            result.append(line)

    cleaned = '\n'.join(result).strip()
    cleaned = re.sub(r'(?<!\w)[^a-zA-Z0-9\xe0-\xff\s.,;:!?\xbf\xa1\-\(\)\[\]"\']{1,3}(?!\w)', '', cleaned)
    cleaned = re.sub(r' {2,}', ' ', cleaned)
    return cleaned.strip()


def ocr_quality_score(text: str) -> dict:
    if not text.strip():
        return {"score": 0.0, "level": "bad", "issues": ["Sin texto detectado"]}
    body = re.sub(r'^---.*?---\s*', '', text, flags=re.DOTALL).strip()
    if not body:
        body = text
    issues = []
    penalties = 0.0
    total_chars = len(body)
    alpha_chars  = sum(1 for c in body if c.isalpha())
    symbol_chars = sum(1 for c in body if not c.isalnum() and c not in ' \n\t.,;:!?¿¡-_()[]{}"\'/\\')
    alpha_ratio  = alpha_chars / max(total_chars, 1)
    if alpha_ratio < 0.35:
        penalties += 0.45
        issues.append(f"Pocos caracteres alfabéticos ({alpha_ratio:.0%})")
    elif alpha_ratio < 0.55:
        penalties += 0.20
        issues.append(f"Proporción alfabética baja ({alpha_ratio:.0%})")
    symbol_ratio = symbol_chars / max(total_chars, 1)
    if symbol_ratio > 0.15:
        penalties += 0.25
        issues.append(f"Alto contenido de símbolos raros ({symbol_ratio:.0%})")
    lines = [l for l in body.splitlines() if l.strip()]
    noisy_lines = 0
    for line in lines:
        alnum = sum(1 for c in line if c.isalnum())
        if len(line) > 3 and alnum / len(line) < 0.45:
            noisy_lines += 1
    noisy_ratio = noisy_lines / max(len(lines), 1)
    if noisy_ratio > 0.40:
        penalties += 0.20
        issues.append(f"Muchas líneas con ruido ({noisy_ratio:.0%} del total)")
    words = body.split()
    short_words = sum(1 for w in words if len(re.sub(r'\W', '', w)) <= 2)
    short_ratio = short_words / max(len(words), 1)
    if short_ratio > 0.55:
        penalties += 0.15
        issues.append(f"Muchas palabras fragmentadas ({short_ratio:.0%})")
    score = max(0.0, 1.0 - penalties)
    level = "ok" if score >= 0.70 else "warning" if score >= 0.40 else "bad"
    return {"score": score, "level": level, "issues": issues}


# ══════════════════════════════════════════════════════════════════
#  HILO IA  —  v4.4: MULTIMODAL (imagen + texto OCR)
# ══════════════════════════════════════════════════════════════════

AI_PROVIDERS = {
    "Claude (Anthropic)": {
        "key_hint": "sk-ant-...",
        "vision": True,
        "desc": (
            "Modelo: claude-sonnet-4 (visión)\n"
            "Recibe imagen + texto OCR para comparar y reconstruir.\n"
            "Costo aprox: $0.004 USD por llamada típica."
        ),
    },
    "Gemini (Google)": {
        "key_hint": "AIzaSy...",
        "vision": True,
        "desc": (
            "Modelo: gemini-1.5-flash (visión)\n"
            "Recibe imagen + texto OCR. Gratis hasta ~1500 req/día.\n"
            "Muy recomendado para uso personal sin costo."
        ),
    },
    "ChatGPT (OpenAI)": {
        "key_hint": "sk-...",
        "vision": True,
        "desc": (
            "Modelo: gpt-4o-mini (visión)\n"
            "Recibe imagen + texto OCR para comparar.\n"
            "Buen balance entre calidad y precio."
        ),
    },
    "DeepSeek": {
        "key_hint": "sk-...",
        "vision": False,
        "desc": (
            "Modelo: deepseek-chat (V3)  —  Solo texto\n"
            "La opción más económica (~10x más barata que Claude).\n"
            "No soporta visión; usa solo el texto OCR."
        ),
    },
}


class AIImproveThread(QThread):
    result_ready   = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self, raw_text: str, api_key: str, provider: str,
                 image_path: str = "", parent=None):
        super().__init__(parent)
        self.raw_text   = raw_text
        self.api_key    = api_key
        self.provider   = provider
        self.image_path = image_path  # v4.4: ruta de la imagen original

    # ── Prompt del sistema ────────────────────────────────────────
    SYSTEM_PROMPT = (
        "Eres un asistente especializado en corregir texto extraído por OCR.\n\n"
        "Se te proporciona:\n"
        "  1. La IMAGEN original del documento.\n"
        "  2. El TEXTO OCR tal como fue extraído (puede tener errores y palabras faltantes).\n\n"
        "Tu tarea:\n"
        "1. Compara la imagen con el texto OCR e identifica las diferencias.\n"
        "2. Reconstruye el texto completo y correcto leyendo directamente de la imagen.\n"
        "3. Corrige errores del OCR: letras confundidas, palabras rotas, texto omitido.\n"
        "4. Mantén la estructura original: párrafos, listas, tablas, jerarquía de títulos.\n"
        "5. NO inventes contenido que no esté en la imagen.\n"
        "6. Si una zona es ilegible, usa [?].\n"
        "7. Responde SOLO con el texto corregido, sin explicaciones ni comentarios."
    )

    # ── Prompt para modo solo-texto (DeepSeek / sin imagen) ───────
    SYSTEM_PROMPT_TEXT_ONLY = (
        "Eres un asistente especializado en limpiar y reconstruir texto "
        "extraído por OCR de imágenes. El texto puede tener errores, "
        "palabras partidas, caracteres basura y frases incoherentes.\n\n"
        "Tu tarea:\n"
        "1. Reconstruye el texto original lo más fielmente posible.\n"
        "2. Corrige solo errores claros del OCR (letras confundidas, palabras "
        "   rotas, símbolos en lugar de letras).\n"
        "3. NO añadas contenido que no estuviera en el original.\n"
        "4. Mantén la estructura (párrafos, listas, bloques de código si los hay).\n"
        "5. Si hay fragmentos completamente irreconocibles, márcalos con [?].\n"
        "6. Responde SOLO con el texto corregido, sin explicaciones ni comentarios."
    )

    def run(self):
        fm_match = re.match(r'^(---.*?---\s*)', self.raw_text, re.DOTALL)
        frontmatter = fm_match.group(1) if fm_match else ""
        body = self.raw_text[len(frontmatter):].strip()

        provider_info = AI_PROVIDERS.get(self.provider, {})
        supports_vision = provider_info.get("vision", False)

        # Cargar imagen si el proveedor soporta visión y existe la ruta
        img_b64 = None
        img_mime = "image/jpeg"
        if supports_vision and self.image_path and os.path.isfile(self.image_path):
            try:
                pil_img = Image.open(self.image_path)
                img_b64, img_mime = pil_to_base64(pil_img, max_side=1600)
            except Exception as e:
                print(f"[WARN] No se pudo cargar imagen para IA: {e}")

        user_prompt = (
            f"Texto OCR extraído del documento (puede estar incompleto o con errores):\n\n"
            f"{body}"
        )

        try:
            if self.provider == "Claude (Anthropic)":
                result = self._call_claude(user_prompt, img_b64, img_mime)
            elif self.provider == "Gemini (Google)":
                result = self._call_gemini(user_prompt, img_b64, img_mime)
            elif self.provider == "ChatGPT (OpenAI)":
                result = self._call_openai(user_prompt, img_b64, img_mime)
            elif self.provider == "DeepSeek":
                result = self._call_deepseek(user_prompt)
            else:
                raise ValueError(f"Proveedor desconocido: {self.provider}")
            final = frontmatter + result if frontmatter else result
            self.result_ready.emit(final)
        except urllib.error.HTTPError as e:
            body_err = e.read().decode("utf-8", errors="replace")
            self.error_occurred.emit(f"HTTP {e.code}: {body_err[:400]}")
        except Exception as e:
            self.error_occurred.emit(str(e))

    # ── Claude (visión multimodal) ────────────────────────────────
    def _call_claude(self, user_prompt: str, img_b64: str | None, img_mime: str) -> str:
        if img_b64:
            # Mensaje multimodal: imagen + texto
            content = [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": img_mime,
                        "data": img_b64,
                    }
                },
                {"type": "text", "text": user_prompt}
            ]
        else:
            content = user_prompt

        payload = json.dumps({
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 2048,
            "system": self.SYSTEM_PROMPT if img_b64 else self.SYSTEM_PROMPT_TEXT_ONLY,
            "messages": [{"role": "user", "content": content}]
        }).encode("utf-8")

        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={"Content-Type": "application/json",
                     "x-api-key": self.api_key,
                     "anthropic-version": "2023-06-01"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data["content"][0]["text"].strip()

    # ── Gemini (visión multimodal) ────────────────────────────────
    def _call_gemini(self, user_prompt: str, img_b64: str | None, img_mime: str) -> str:
        combined_text = f"{self.SYSTEM_PROMPT if img_b64 else self.SYSTEM_PROMPT_TEXT_ONLY}\n\n{user_prompt}"

        if img_b64:
            parts = [
                {"inline_data": {"mime_type": img_mime, "data": img_b64}},
                {"text": combined_text}
            ]
        else:
            parts = [{"text": combined_text}]

        payload = json.dumps({
            "contents": [{"parts": parts}],
            "generationConfig": {"maxOutputTokens": 2048}
        }).encode("utf-8")

        url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
               f"gemini-1.5-flash:generateContent?key={self.api_key}")
        req = urllib.request.Request(url, data=payload,
                                     headers={"Content-Type": "application/json"},
                                     method="POST")
        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data["candidates"][0]["content"]["parts"][0]["text"].strip()

    # ── OpenAI GPT-4o (visión multimodal) ────────────────────────
    def _call_openai(self, user_prompt: str, img_b64: str | None, img_mime: str) -> str:
        if img_b64:
            user_content = [
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{img_mime};base64,{img_b64}"}
                },
                {"type": "text", "text": user_prompt}
            ]
        else:
            user_content = user_prompt

        payload = json.dumps({
            "model": "gpt-4o-mini",
            "max_tokens": 2048,
            "messages": [
                {"role": "system",
                 "content": self.SYSTEM_PROMPT if img_b64 else self.SYSTEM_PROMPT_TEXT_ONLY},
                {"role": "user", "content": user_content}
            ]
        }).encode("utf-8")

        req = urllib.request.Request(
            "https://api.openai.com/v1/chat/completions",
            data=payload,
            headers={"Content-Type": "application/json",
                     "Authorization": f"Bearer {self.api_key}"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data["choices"][0]["message"]["content"].strip()

    # ── DeepSeek (solo texto) ─────────────────────────────────────
    def _call_deepseek(self, user_prompt: str) -> str:
        payload = json.dumps({
            "model": "deepseek-chat",
            "max_tokens": 2048,
            "messages": [
                {"role": "system", "content": self.SYSTEM_PROMPT_TEXT_ONLY},
                {"role": "user", "content": user_prompt}
            ]
        }).encode("utf-8")
        req = urllib.request.Request(
            "https://api.deepseek.com/v1/chat/completions",
            data=payload,
            headers={"Content-Type": "application/json",
                     "Authorization": f"Bearer {self.api_key}"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data["choices"][0]["message"]["content"].strip()


# ══════════════════════════════════════════════════════════════════
#  FRONTMATTER YAML
# ══════════════════════════════════════════════════════════════════

def generate_frontmatter(title: str = "Nota Importada", tags: list | None = None) -> str:
    fecha = datetime.now().strftime("%Y-%m-%d")
    hora  = datetime.now().strftime("%H:%M")
    tag_str = ", ".join(tags) if tags else "ocr, importado"
    return (
        f"---\n"
        f"title: {title}\n"
        f"date: {fecha}\n"
        f"time: {hora}\n"
        f"tags: [{tag_str}]\n"
        f"source: OCR - Syntax Vision v4.4\n"
        f"---\n\n"
    )


# ══════════════════════════════════════════════════════════════════
#  EXPORTADORES
# ══════════════════════════════════════════════════════════════════

def export_to_pdf(markdown_text: str, output_path: str) -> tuple:
    try:
        doc = SimpleDocTemplate(
            output_path, pagesize=A4,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            topMargin=2.5*cm,  bottomMargin=2.5*cm,
        )
        styles = getSampleStyleSheet()
        F = "Helvetica"
        s_body = ParagraphStyle("SV_Body", parent=styles["Normal"],
                                fontName=F, fontSize=10, leading=16, spaceAfter=4)
        s_h1   = ParagraphStyle("SV_H1", parent=styles["Heading1"],
                                fontName=F+"-Bold", fontSize=17, spaceAfter=8, spaceBefore=14,
                                textColor=colors.HexColor("#11212D"))
        s_h2   = ParagraphStyle("SV_H2", parent=styles["Heading2"],
                                fontName=F+"-Bold", fontSize=13, spaceAfter=6, spaceBefore=10,
                                textColor=colors.HexColor("#253745"))
        s_h3   = ParagraphStyle("SV_H3", parent=styles["Heading3"],
                                fontName=F+"-Bold", fontSize=11, spaceAfter=4, spaceBefore=8,
                                textColor=colors.HexColor("#4A5C6A"))
        s_code = ParagraphStyle("SV_Code", parent=styles["Code"],
                                fontName="Courier", fontSize=9, leading=14,
                                backColor=colors.HexColor("#F4F4F4"),
                                leftIndent=20, spaceAfter=6)
        s_meta = ParagraphStyle("SV_Meta", parent=styles["Normal"],
                                fontName=F, fontSize=8,
                                textColor=colors.HexColor("#9BA8AB"), spaceAfter=2)
        story = []
        lines = markdown_text.split('\n')
        in_fm = False
        i = 0
        while i < len(lines):
            raw_line = lines[i]
            if i == 0 and raw_line.strip() == '---':
                in_fm = True; i += 1; continue
            if in_fm:
                if raw_line.strip() == '---':
                    in_fm = False
                    story.append(HRFlowable(width="100%", thickness=0.5,
                                            color=colors.HexColor("#9BA8AB")))
                    story.append(Spacer(1, 0.3*cm))
                else:
                    safe = raw_line.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                    story.append(Paragraph(safe, s_meta))
                i += 1; continue
            if raw_line.startswith('### '):
                story.append(Paragraph(raw_line[4:], s_h3))
            elif raw_line.startswith('## '):
                story.append(Paragraph(raw_line[3:], s_h2))
            elif raw_line.startswith('# '):
                story.append(Paragraph(raw_line[2:], s_h1))
            elif raw_line.strip() == '':
                story.append(Spacer(1, 0.18*cm))
            elif raw_line.startswith('    ') or raw_line.startswith('\t'):
                safe = raw_line.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                story.append(Paragraph(safe, s_code))
            else:
                safe = raw_line.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                safe = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', safe)
                safe = re.sub(r'\*(.*?)\*',     r'<i>\1</i>', safe)
                safe = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', safe)
                try:
                    story.append(Paragraph(safe, s_body))
                except Exception:
                    story.append(Paragraph(raw_line, s_body))
            i += 1
        if not story:
            story.append(Paragraph("(documento vacío)", s_body))
        doc.build(story)
        return True, ""
    except Exception as e:
        return False, str(e)


def export_to_docx(markdown_text: str, output_path: str) -> tuple:
    try:
        doc = DocxDocument()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)
        lines = markdown_text.split('\n')
        in_fm = False
        i = 0
        while i < len(lines):
            line = lines[i]
            if i == 0 and line.strip() == '---':
                in_fm = True; i += 1; continue
            if in_fm:
                if line.strip() == '---':
                    in_fm = False
                    p = doc.add_paragraph('─' * 55)
                    p.runs[0].font.color.rgb = RGBColor(0x9B, 0xA8, 0xAB)
                else:
                    p = doc.add_paragraph(line)
                    if p.runs:
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.color.rgb = RGBColor(0x9B, 0xA8, 0xAB)
                i += 1; continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.strip() == '':
                doc.add_paragraph('')
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
                for part in parts:
                    run = p.add_run()
                    if part.startswith('**') and part.endswith('**'):
                        run.text = part[2:-2]; run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run.text = part[1:-1]; run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        run.text = part[1:-1]
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        run.text = part
            i += 1
        doc.save(output_path)
        return True, ""
    except Exception as e:
        return False, str(e)


def export_to_txt(text: str, path: str) -> tuple:
    try:
        with open(path, 'w', encoding='utf-8') as f:
            f.write(text)
        return True, ""
    except Exception as e:
        return False, str(e)


# ══════════════════════════════════════════════════════════════════
#  HILO DE CÁMARA
# ══════════════════════════════════════════════════════════════════

class CameraThread(QThread):
    frame_ready  = pyqtSignal(QImage)
    camera_error = pyqtSignal(str)

    def __init__(self, camera_index: int = 0):
        super().__init__()
        self.camera_index = camera_index
        self._running = False
        self.cap = None

    def run(self):
        backend = cv2.CAP_DSHOW if os.name == "nt" else cv2.CAP_ANY
        self.cap = cv2.VideoCapture(self.camera_index, backend)
        if not self.cap.isOpened():
            self.camera_error.emit(
                f"No se pudo abrir la cámara {self.camera_index}.\n"
                "Verifica que no esté en uso por otra aplicación."
            )
            return
        self._running = True
        while self._running:
            ret, frame = self.cap.read()
            if not ret:
                break
            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            h, w, ch = rgb.shape
            qi = QImage(rgb.data, w, h, ch * w, QImage.Format.Format_RGB888)
            self.frame_ready.emit(qi.copy())
            self.msleep(33)
        if self.cap:
            self.cap.release()

    def capture_frame(self) -> QImage | None:
        if self.cap and self.cap.isOpened():
            ret, frame = self.cap.read()
            if ret:
                rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                h, w, ch = rgb.shape
                return QImage(rgb.data, w, h, ch * w, QImage.Format.Format_RGB888).copy()
        return None

    def stop(self):
        self._running = False
        self.wait(2500)
        if self.cap:
            self.cap.release()
            self.cap = None


# ══════════════════════════════════════════════════════════════════
#  WIDGET DRAG & DROP
# ══════════════════════════════════════════════════════════════════

class DropArea(QLabel):
    image_dropped = pyqtSignal(str)
    VALID_EXT = ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp')

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("DropArea")
        self.setAcceptDrops(True)
        self.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.setWordWrap(True)
        self._reset_text()

    def _reset_text(self):
        self.setText("⬇  Arrastra una imagen aquí\n\nPNG · JPG · BMP · TIFF")
        self.setPixmap(QPixmap())

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if not urls:
            return
        path = urls[0].toLocalFile()
        if path.lower().endswith(self.VALID_EXT):
            self.image_dropped.emit(path)
        else:
            QMessageBox.warning(self, "Formato no soportado",
                                f"Usa: {', '.join(self.VALID_EXT)}")

    def show_preview(self, pixmap: QPixmap):
        self.setPixmap(pixmap.scaled(
            self.size(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        ))
        self.setText("")


# ══════════════════════════════════════════════════════════════════
#  DIÁLOGO DE CÁMARA
# ══════════════════════════════════════════════════════════════════

class CameraDialog(QDialog):
    photo_taken = pyqtSignal(QImage)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Captura con Cámara")
        self.setFixedSize(740, 640)
        self.setModal(True)
        self.cam_thread: CameraThread | None = None
        self._cameras: list[dict] = []
        self._build_ui()
        self._detect_and_populate()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)
        cam_row = QHBoxLayout()
        lbl_cam = QLabel("CÁMARA:")
        lbl_cam.setFixedWidth(70)
        cam_row.addWidget(lbl_cam)
        self.cam_combo = QComboBox()
        self.cam_combo.setMinimumWidth(360)
        self.cam_combo.currentIndexChanged.connect(self._on_camera_changed)
        cam_row.addWidget(self.cam_combo, stretch=1)
        self.btn_refresh = QPushButton("🔄")
        self.btn_refresh.setFixedWidth(38)
        self.btn_refresh.setToolTip("Volver a detectar cámaras")
        self.btn_refresh.clicked.connect(self._detect_and_populate)
        cam_row.addWidget(self.btn_refresh)
        layout.addLayout(cam_row)
        self.lbl_cam_status = QLabel("Detectando cámaras…")
        self.lbl_cam_status.setStyleSheet("font-size: 10px; font-weight: 400; letter-spacing: 0;")
        layout.addWidget(self.lbl_cam_status)
        self.view = QLabel()
        self.view.setObjectName("CameraView")
        self.view.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.view.setMinimumSize(660, 495)
        self.view.setText("⏳  Iniciando cámara…")
        layout.addWidget(self.view, stretch=1)
        btn_row = QHBoxLayout()
        self.btn_snap = QPushButton("📷  Tomar Foto")
        self.btn_snap.setObjectName("AccentButton")
        self.btn_snap.setEnabled(False)
        self.btn_snap.clicked.connect(self._take_photo)
        btn_close = QPushButton("✕  Cerrar")
        btn_close.setObjectName("DangerButton")
        btn_close.clicked.connect(self._close_cam)
        btn_row.addWidget(self.btn_snap)
        btn_row.addWidget(btn_close)
        layout.addLayout(btn_row)

    def _detect_and_populate(self):
        self._stop_thread()
        self.btn_snap.setEnabled(False)
        self.btn_refresh.setEnabled(False)
        self.cam_combo.blockSignals(True)
        self.cam_combo.clear()
        self.cam_combo.addItem("🔍  Detectando cámaras…")
        self.lbl_cam_status.setText("Buscando dispositivos, un momento…")
        QApplication.processEvents()
        self._detect_thread = CameraDetectThread()
        self._detect_thread.cameras_found.connect(self._on_cameras_detected)
        self._detect_thread.start()

    def _on_cameras_detected(self, cameras: list):
        self._cameras = cameras
        self.cam_combo.blockSignals(True)
        self.cam_combo.clear()
        for cam in cameras:
            self.cam_combo.addItem(f"📷  {cam['label']}", userData=cam['index'])
        self.cam_combo.blockSignals(False)
        self.btn_refresh.setEnabled(True)
        n = len(cameras)
        self.lbl_cam_status.setText(
            f"✔  {n} cámara{'s' if n != 1 else ''} encontrada{'s' if n != 1 else ''}"
            "  —  Selecciona y presiona Tomar Foto"
        )
        if cameras:
            self._start_camera(cameras[0]['index'])

    def _on_camera_changed(self, idx: int):
        if idx < 0 or idx >= len(self._cameras):
            return
        self._start_camera(self._cameras[idx]['index'])

    def _start_camera(self, cam_index: int):
        self._stop_thread()
        self.btn_snap.setEnabled(False)
        self.view.setText(f"⏳  Iniciando cámara {cam_index}…")
        QApplication.processEvents()
        self.cam_thread = CameraThread(cam_index)
        self.cam_thread.frame_ready.connect(self._update_frame)
        self.cam_thread.camera_error.connect(self._handle_error)
        self.cam_thread.start()

    def _update_frame(self, qi: QImage):
        self.btn_snap.setEnabled(True)
        self.view.setPixmap(QPixmap.fromImage(qi).scaled(
            self.view.size(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        ))

    def _take_photo(self):
        if self.cam_thread is None:
            return
        img = self.cam_thread.capture_frame()
        if img:
            self.photo_taken.emit(img)
            self._close_cam()
        else:
            QMessageBox.warning(self, "Error", "No se pudo capturar el frame.")

    def _handle_error(self, msg: str):
        self.btn_snap.setEnabled(False)
        self.view.setText(f"⚠  {msg}")
        self.lbl_cam_status.setText("✗  Error al abrir la cámara seleccionada")

    def _stop_thread(self):
        if self.cam_thread is not None:
            self.cam_thread.stop()
            self.cam_thread = None

    def _close_cam(self):
        self._stop_thread()
        self.close()

    def closeEvent(self, event):
        self._stop_thread()
        super().closeEvent(event)


# ══════════════════════════════════════════════════════════════════
#  PANEL DE TEMAS
# ══════════════════════════════════════════════════════════════════

class ThemePanel(QWidget):
    theme_changed = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(6)
        lbl = QLabel("TEMA DE INTERFAZ")
        lbl.setObjectName("SectionTitle")
        layout.addWidget(lbl)
        self.combo = QComboBox()
        self.combo.addItems(list(THEMES.keys()))
        self.combo.setCurrentText(CURRENT_THEME)
        self.combo.currentTextChanged.connect(self.theme_changed.emit)
        layout.addWidget(self.combo)
        self.preview = QLabel()
        self.preview.setFixedHeight(28)
        self.preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.preview.setStyleSheet("border-radius: 5px; font-size: 10px;")
        layout.addWidget(self.preview)
        self._update_preview(CURRENT_THEME)
        self.combo.currentTextChanged.connect(self._update_preview)

    def _update_preview(self, name: str):
        if name not in THEMES:
            return
        t = THEMES[name]
        swatches = [t['bg_root'], t['bg_panel'], t['bg_widget'],
                    t['accent'], t['highlight'], t['text_main']]
        blocks = "".join(
            f'<span style="background:{c};padding:0 10px;">&nbsp;</span>'
            for c in swatches
        )
        self.preview.setText(blocks)
        self.preview.setStyleSheet(
            f"background: {t['bg_panel']}; border-radius: 5px;"
            f"border: 1px solid {t['border']}; padding: 2px;"
        )

    def set_active(self, name: str):
        self.combo.setCurrentText(name)


# ══════════════════════════════════════════════════════════════════
#  DIÁLOGO DE BIENVENIDA
# ══════════════════════════════════════════════════════════════════

class WelcomeDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Syntax Vision — Aviso importante")
        self.setFixedSize(520, 420)
        self.setModal(True)
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(14)
        title = QLabel("⚠  ANTES DE COMENZAR")
        title.setObjectName("AppTitle")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        layout.addWidget(sep)
        warnings = [
            ("📷  CALIDAD DE IMAGEN",
             "Asegúrate de que la imagen sea nítida, bien iluminada y sin sombras. "
             "Imágenes borrosas, con brillo excesivo o muy oscuras producirán "
             "resultados deficientes."),
            ("🔤  SOLO LETRAS MOLDE O IMPRESAS",
             "El OCR reconoce texto impreso, mecanografiado y letra de molde clara. "
             "La letra cursiva o manuscrita libre tiene reconocimiento limitado — "
             "usa el agente EasyOCR y el modo 'photo' para mejores resultados."),
            ("✏️  VERIFICA EL RESULTADO",
             "El OCR está sujeto a errores. Siempre revisa y corrige el texto "
             "en el editor antes de exportar. Usa '✨ Mejorar con IA' para "
             "reconstruir texto dañado automáticamente con visión IA."),
        ]
        for icon_title, body in warnings:
            box = QGroupBox(icon_title)
            box_layout = QVBoxLayout(box)
            lbl = QLabel(body)
            lbl.setWordWrap(True)
            lbl.setStyleSheet("font-size: 12px; font-weight: 400; letter-spacing: 0;")
            box_layout.addWidget(lbl)
            layout.addWidget(box)
        layout.addStretch()
        btn_ok = QPushButton("✔  Entendido, comenzar")
        btn_ok.setObjectName("AccentButton")
        btn_ok.setFixedHeight(38)
        btn_ok.clicked.connect(self.accept)
        layout.addWidget(btn_ok)


# ══════════════════════════════════════════════════════════════════
#  EDITOR CON LÍMITE DE CARACTERES
# ══════════════════════════════════════════════════════════════════

class LimitedTextEdit(QTextEdit):
    char_count_changed = pyqtSignal(int)

    def __init__(self, max_chars: int = MAX_CHARS, parent=None):
        super().__init__(parent)
        self.max_chars = max_chars
        self.textChanged.connect(self._enforce_limit)

    def _enforce_limit(self):
        text = self.toPlainText()
        count = len(text)
        if count > self.max_chars:
            cursor = self.textCursor()
            pos = cursor.position()
            self.blockSignals(True)
            self.setPlainText(text[:self.max_chars])
            new_pos = min(pos, self.max_chars)
            cursor.setPosition(new_pos)
            self.setTextCursor(cursor)
            self.blockSignals(False)
            count = self.max_chars
        self.char_count_changed.emit(count)


# ══════════════════════════════════════════════════════════════════
#  VENTANA PRINCIPAL
# ══════════════════════════════════════════════════════════════════

class SyntaxVision(QMainWindow):

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Syntax Vision  ·  OCR → PKM  ·  v4.4")
        self.resize(1500, 900)
        self.setMinimumSize(1100, 700)
        if os.path.isfile(LOGO_PATH):
            self.setWindowIcon(QIcon(LOGO_PATH))
        else:
            print(f"[INFO] Logo no encontrado en: {LOGO_PATH}")
        self._current_image_path: str = ""
        self._current_ocr_mode: str = "auto"
        self._current_engine: str = "tesseract"
        self._api_keys: dict = {}
        self._db = db_init()
        self._build_ui()
        self._apply_theme(CURRENT_THEME)
        self._refresh_history_widget()
        dlg = WelcomeDialog(self)
        dlg.exec()

    def _build_ui(self):
        root = QWidget()
        root_layout = QVBoxLayout(root)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)
        self.setCentralWidget(root)
        root_layout.addWidget(self._build_header())
        body = QWidget()
        body_layout = QHBoxLayout(body)
        body_layout.setContentsMargins(10, 8, 10, 10)
        body_layout.setSpacing(8)
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.addWidget(self._build_capture_panel())
        splitter.addWidget(self._build_editor_panel())
        splitter.addWidget(self._build_right_panel())
        splitter.setSizes([320, 740, 280])
        body_layout.addWidget(splitter)
        root_layout.addWidget(body, stretch=1)

    def _build_header(self) -> QWidget:
        w = QWidget()
        w.setObjectName("Header")
        layout = QHBoxLayout(w)
        layout.setContentsMargins(18, 10, 18, 10)
        if os.path.isfile(LOGO_PATH):
            logo_lbl = QLabel()
            pixmap = QPixmap(LOGO_PATH).scaled(36, 36,
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation)
            logo_lbl.setPixmap(pixmap)
            logo_lbl.setFixedSize(40, 40)
            logo_lbl.setStyleSheet("background: transparent;")
            layout.addWidget(logo_lbl)
        title = QLabel("SYNTAX VISION")
        title.setObjectName("AppTitle")
        subtitle = QLabel("OCR  ·  PKM  ·  MARKDOWN  ·  v4.4")
        subtitle.setObjectName("AppSubtitle")
        vbox = QVBoxLayout()
        vbox.setSpacing(0)
        vbox.addWidget(title)
        vbox.addWidget(subtitle)
        layout.addLayout(vbox)
        layout.addStretch()
        self.lbl_status = QLabel("en espera")
        self.lbl_status.setObjectName("AppSubtitle")
        layout.addWidget(self.lbl_status)
        return w

    def _build_capture_panel(self) -> QWidget:
        w = QWidget()
        w.setMinimumWidth(270)
        layout = QVBoxLayout(w)
        layout.setSpacing(8)
        lbl = QLabel("CAPTURA")
        lbl.setObjectName("SectionTitle")
        layout.addWidget(lbl)
        self.drop_area = DropArea()
        self.drop_area.setMinimumHeight(190)
        self.drop_area.image_dropped.connect(self._load_image_from_path)
        layout.addWidget(self.drop_area)
        btn_file = QPushButton("📂  Abrir Imagen")
        btn_file.clicked.connect(self._open_file_dialog)
        layout.addWidget(btn_file)
        btn_cam = QPushButton("📷  Cámara Web")
        btn_cam.setObjectName("AccentButton")
        btn_cam.clicked.connect(self._open_camera)
        layout.addWidget(btn_cam)

        agent_group = QGroupBox("AGENTE OCR")
        agent_layout = QVBoxLayout(agent_group)
        agent_layout.setSpacing(6)
        self.agent_combo = QComboBox()
        self.agent_combo.addItems([
            "🖨  Tesseract  —  Texto impreso / letra molde",
            "✍  EasyOCR   —  Manuscrito / mixto",
        ])
        self.agent_combo.currentIndexChanged.connect(self._update_agent)
        agent_layout.addWidget(self.agent_combo)
        self.agent_hint = QLabel(
            "Rápido. Ideal para documentos impresos\n"
            "y texto en letra molde."
        )
        self.agent_hint.setWordWrap(True)
        agent_layout.addWidget(self.agent_hint)
        self.agent_badge = QLabel("● Listo")
        self.agent_badge.setStyleSheet("color: #4caf50; font-size: 10px; font-weight: 700;")
        agent_layout.addWidget(self.agent_badge)
        layout.addWidget(agent_group)

        ocr_group = QGroupBox("PREPROCESAMIENTO")
        ocr_layout = QVBoxLayout(ocr_group)
        ocr_layout.setSpacing(4)
        self.ocr_mode_combo = QComboBox()
        self.ocr_mode_combo.addItems([
            "auto  — Detección automática",
            "photo  — Foto de papel / cuaderno",
            "document — Documento oficial (constancias, diplomas, formularios)",
            "raw    — Sin preprocesamiento",
        ])
        self.ocr_mode_combo.currentIndexChanged.connect(self._update_ocr_mode)
        ocr_layout.addWidget(self.ocr_mode_combo)
        self.preproc_hint = QLabel(
            "'photo' para fotos de papel o cuadernos\n"
            "'document' para constancias, sellos y formularios"
        )
        self.preproc_hint.setWordWrap(True)
        ocr_layout.addWidget(self.preproc_hint)
        layout.addWidget(ocr_group)

        lang_group = QGroupBox("IDIOMA")
        lang_layout = QHBoxLayout(lang_group)
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["spa", "eng", "spa+eng", "fra", "deu", "ita", "por"])
        lang_layout.addWidget(self.lang_combo)
        layout.addWidget(lang_group)

        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        layout.addWidget(sep)

        self.btn_ocr = QPushButton("🔍  Ejecutar OCR")
        self.btn_ocr.setObjectName("AccentButton")
        self.btn_ocr.setEnabled(False)
        self.btn_ocr.clicked.connect(self._run_ocr)
        layout.addWidget(self.btn_ocr)
        layout.addStretch()
        return w

    def _build_editor_panel(self) -> QWidget:
        w = QWidget()
        layout = QVBoxLayout(w)
        layout.setSpacing(8)
        header_row = QHBoxLayout()
        lbl = QLabel("EDITOR  —  MARKDOWN / OCR")
        lbl.setObjectName("SectionTitle")
        header_row.addWidget(lbl)
        header_row.addStretch()
        self.ai_provider_combo = QComboBox()
        self.ai_provider_combo.addItems(list(AI_PROVIDERS.keys()))
        self.ai_provider_combo.setFixedWidth(200)
        self.ai_provider_combo.currentTextChanged.connect(self._update_ai_provider_hint)
        header_row.addWidget(self.ai_provider_combo)
        self.btn_ai = QPushButton("✨  Mejorar con IA")
        self.btn_ai.setObjectName("AccentButton")
        self.btn_ai.setEnabled(False)
        self.btn_ai.setToolTip("Envía imagen + texto OCR a la IA para reconstruir el texto completo")
        self.btn_ai.clicked.connect(self._run_ai_improve)
        header_row.addWidget(self.btn_ai)
        layout.addLayout(header_row)
        self.ai_provider_hint = QLabel()
        self.ai_provider_hint.setWordWrap(True)
        self.ai_provider_hint.setStyleSheet("font-size: 10px; font-weight: 400; letter-spacing: 0;")
        layout.addWidget(self.ai_provider_hint)
        self._update_ai_provider_hint(self.ai_provider_combo.currentText())
        self.quality_banner = QLabel()
        self.quality_banner.setWordWrap(True)
        self.quality_banner.setVisible(False)
        self.quality_banner.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
        self.quality_banner.setObjectName("QualityBanner")
        self.quality_banner.setMinimumHeight(36)
        self.quality_banner.setContentsMargins(10, 6, 10, 6)
        layout.addWidget(self.quality_banner)
        self.editor = LimitedTextEdit(max_chars=MAX_CHARS)
        self.editor.setPlaceholderText(
            "El resultado del OCR aparecerá aquí.\n"
            "Edita el texto antes de exportar.\n\n"
            "Tip: usa 'document' para constancias y formularios.\n"
            "Luego usa '✨ Mejorar con IA' — envía la imagen original\n"
            "para que la IA compare y recupere el texto faltante."
        )
        self.editor.char_count_changed.connect(self._update_char_counter)
        layout.addWidget(self.editor, stretch=1)
        counter_row = QHBoxLayout()
        counter_row.addStretch()
        self.lbl_char_counter = QLabel(f"0 / {MAX_CHARS:,}")
        self.lbl_char_counter.setObjectName("CharCounter")
        self._style_char_counter(0)
        counter_row.addWidget(self.lbl_char_counter)
        layout.addLayout(counter_row)
        exp_group = QGroupBox("EXPORTAR COMO")
        exp_layout = QHBoxLayout(exp_group)
        exp_layout.setSpacing(6)
        self.btn_md   = QPushButton("💾  .md")
        self.btn_txt  = QPushButton("📄  .txt")
        self.btn_docx = QPushButton("📝  .docx")
        self.btn_pdf  = QPushButton("🖨  .pdf")
        self.btn_docx.setObjectName("AccentButton")
        self.btn_pdf.setObjectName("AccentButton")
        for btn in (self.btn_md, self.btn_txt, self.btn_docx, self.btn_pdf):
            btn.setEnabled(False)
        self.btn_md.clicked.connect(lambda: self._export("md"))
        self.btn_txt.clicked.connect(lambda: self._export("txt"))
        self.btn_docx.clicked.connect(lambda: self._export("docx"))
        self.btn_pdf.clicked.connect(lambda: self._export("pdf"))
        self.btn_clear = QPushButton("🗑")
        self.btn_clear.setObjectName("DangerButton")
        self.btn_clear.setFixedWidth(44)
        self.btn_clear.setToolTip("Limpiar editor")
        self.btn_clear.clicked.connect(self._clear_editor)
        for btn in (self.btn_md, self.btn_txt, self.btn_docx, self.btn_pdf):
            exp_layout.addWidget(btn)
        exp_layout.addStretch()
        exp_layout.addWidget(self.btn_clear)
        layout.addWidget(exp_group)
        return w

    def _build_right_panel(self) -> QWidget:
        w = QWidget()
        w.setMinimumWidth(220)
        w.setMaximumWidth(320)
        layout = QVBoxLayout(w)
        layout.setSpacing(10)
        self.theme_panel = ThemePanel()
        self.theme_panel.theme_changed.connect(self._apply_theme)
        layout.addWidget(self.theme_panel)
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        layout.addWidget(sep)
        hist_lbl = QLabel(f"HISTORIAL PERSISTENTE (máx. {MAX_HISTORY})")
        hist_lbl.setObjectName("SectionTitle")
        layout.addWidget(hist_lbl)
        self.history_list = QListWidget()
        self.history_list.itemClicked.connect(self._load_from_history)
        layout.addWidget(self.history_list, stretch=1)
        btn_clear_hist = QPushButton("🗑  Limpiar historial")
        btn_clear_hist.setObjectName("DangerButton")
        btn_clear_hist.clicked.connect(self._clear_history)
        layout.addWidget(btn_clear_hist)
        return w

    def _style_char_counter(self, count: int):
        ratio = count / MAX_CHARS
        if ratio < 0.75:
            color, bg = "#9BA8AB", "transparent"
        elif ratio < 0.90:
            color, bg = "#f0b040", "#2a2000"
        else:
            color, bg = "#f08080", "#2a0000"
        self.lbl_char_counter.setStyleSheet(
            f"QLabel#CharCounter {{ color: {color}; background: {bg};"
            f" border-radius: 4px; padding: 2px 8px; font-size: 11px; font-weight: 700; }}"
        )

    def _update_char_counter(self, count: int):
        self.lbl_char_counter.setText(f"{count:,} / {MAX_CHARS:,}")
        self._style_char_counter(count)

    def _apply_theme(self, name: str):
        global CURRENT_THEME
        CURRENT_THEME = name
        self.setStyleSheet(build_stylesheet(name))
        t = THEMES[name]
        for child in self.findChildren(QWidget):
            if child.objectName() == "Header":
                child.setStyleSheet(
                    f"QWidget#Header {{ background-color: {t['bg_panel']};"
                    f"border-bottom: 1px solid {t['border']}; }}"
                )

    def _open_file_dialog(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Seleccionar Imagen", "",
            "Imágenes (*.png *.jpg *.jpeg *.bmp *.tiff *.webp)"
        )
        if path:
            self._load_image_from_path(path)

    def _load_image_from_path(self, path: str):
        if not os.path.isfile(path):
            QMessageBox.critical(self, "Error", f"Archivo no encontrado:\n{path}")
            return
        pixmap = QPixmap(path)
        if pixmap.isNull():
            QMessageBox.critical(self, "Error", "No se pudo cargar la imagen.")
            return
        self._current_image_path = path
        self.drop_area.show_preview(pixmap)
        self.lbl_status.setText(f"✔ {os.path.basename(path)}")
        self.btn_ocr.setEnabled(True)
        _preproc_cache["path"] = None

    def _open_camera(self):
        dlg = CameraDialog(self)
        dlg.photo_taken.connect(self._load_image_from_qimage)
        dlg.exec()

    def _load_image_from_qimage(self, qi: QImage):
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        tmp = os.path.join(os.path.expanduser("~"), f"sv_cam_{ts}.png")
        qi.save(tmp, "PNG")
        self._load_image_from_path(tmp)

    def _update_ocr_mode(self, idx: int):
        self._current_ocr_mode = ["auto", "photo", "document", "raw"][idx]
        _preproc_cache["mode"] = None

    def _update_agent(self, idx: int):
        self._current_engine = ["tesseract", "easyocr"][idx]
        if self._current_engine == "tesseract":
            self.agent_hint.setText(
                "Rápido. Ideal para documentos impresos\n"
                "y texto en letra molde."
            )
            self.agent_badge.setText("● Listo")
            self.agent_badge.setStyleSheet("color: #4caf50; font-size: 10px; font-weight: 700;")
        else:
            self.agent_hint.setText(
                "Deep-learning. Mejor para letra manuscrita,\n"
                "texto mixto y fotos de cuadernos.\n"
                "⚠ Descarga modelos ~100 MB la 1ª vez."
            )
            self.agent_badge.setText("● Listo (carga en primer uso)")
            self.agent_badge.setStyleSheet("color: #f0b040; font-size: 10px; font-weight: 700;")

    def _run_ocr(self):
        if not self._current_image_path:
            QMessageBox.warning(self, "Sin imagen", "Primero carga una imagen.")
            return
        engine = self._current_engine
        if engine == "easyocr" and _easyocr_reader is None:
            reply = QMessageBox.information(
                self, "EasyOCR — Primera carga",
                "La primera vez EasyOCR descarga modelos (~100 MB).\n"
                "Esto puede tardar 1-3 minutos.\n\n"
                "Las siguientes ejecuciones serán inmediatas.",
                QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel
            )
            if reply == QMessageBox.StandardButton.Cancel:
                return
            self.agent_badge.setText("⏳ Descargando modelos…")
            self.agent_badge.setStyleSheet("color: #f0b040; font-size: 10px; font-weight: 700;")

        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        engine_label = "EasyOCR" if engine == "easyocr" else "Tesseract"
        self.lbl_status.setText(f"⏳ {engine_label} procesando…")
        QApplication.processEvents()

        try:
            pil_img = Image.open(self._current_image_path)
            lang    = self.lang_combo.currentText()
            processed = get_preprocessed(pil_img, self._current_image_path,
                                          self._current_ocr_mode)

            use_zone_ocr = (
                engine == "tesseract"
                and self._current_ocr_mode in ("document", "auto")
            )
            if use_zone_ocr:
                reordered = reorder_ocr_text_by_layout(pil_img, lang=lang)
                if len(reordered.strip()) >= 30:
                    raw_text = reordered
                else:
                    raw_text = run_ocr(processed, engine=engine, lang=lang,
                                       mode=self._current_ocr_mode)
            else:
                raw_text = run_ocr(processed, engine=engine, lang=lang,
                                   mode=self._current_ocr_mode)

        except Exception as e:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(self, f"Error {engine_label}", f"Fallo al procesar:\n{e}")
            self.lbl_status.setText("✗ Error OCR")
            self.agent_badge.setText("✗ Error")
            self.agent_badge.setStyleSheet("color: #f08080; font-size: 10px; font-weight: 700;")
            return
        finally:
            QApplication.restoreOverrideCursor()

        if engine == "easyocr":
            self.agent_badge.setText("● Listo")
            self.agent_badge.setStyleSheet("color: #4caf50; font-size: 10px; font-weight: 700;")

        if len(raw_text.strip()) < 5:
            QMessageBox.warning(
                self, "Sin texto detectado",
                "No se encontró texto legible.\n\n"
                "Sugerencias:\n"
                "• Usa EasyOCR para letra manuscrita o mixta\n"
                "• Cambia el modo a 'document' para constancias\n"
                "• Verifica resolución y contraste de la imagen"
            )
            self.lbl_status.setText("⚠ Sin resultados")
            return

        cleaned = clean_ocr_text(raw_text)

        title, ok = QInputDialog.getText(
            self, "Título de la nota", "Título:",
            QLineEdit.EchoMode.Normal,
            f"Nota {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        if not ok or not title.strip():
            title = "Nota Importada"

        tags = ["ocr", engine_label.lower(), "importado"]
        full_content = generate_frontmatter(title=title.strip(), tags=tags) + cleaned

        if len(full_content) > MAX_CHARS:
            full_content = full_content[:MAX_CHARS]
            QMessageBox.warning(
                self, "Texto truncado",
                f"El texto supera el límite de {MAX_CHARS:,} caracteres y fue truncado.\n"
                "Puedes exportar por secciones si necesitas el texto completo."
            )

        self.editor.setPlainText(full_content)
        for btn in (self.btn_md, self.btn_txt, self.btn_docx, self.btn_pdf):
            btn.setEnabled(True)
        self.btn_ai.setEnabled(True)

        quality = ocr_quality_score(full_content)
        self._show_quality_banner(quality)
        self.lbl_status.setText(f"✔ {engine_label}: {len(cleaned)} chars")
        self._add_to_history(title.strip(), full_content)

    def _show_quality_banner(self, quality: dict):
        level  = quality["level"]
        score  = quality["score"]
        issues = quality["issues"]
        if level == "ok":
            self.quality_banner.setVisible(False)
            return
        if level == "warning":
            bg, border, icon = "#2a2000", "#a07000", "⚠️"
            msg_prefix = "Calidad OCR moderada"
        else:
            bg, border, icon = "#2a0000", "#a03030", "🔴"
            msg_prefix = "Calidad OCR baja — se recomienda revisar el contenido"
        issues_str = "  ·  ".join(issues) if issues else ""
        score_pct  = f"{score:.0%}"
        self.quality_banner.setText(
            f"{icon}  {msg_prefix}  ({score_pct})   "
            + (f"→  {issues_str}" if issues_str else "")
            + "   —   Usa '✨ Mejorar con IA' — la IA verá la imagen original."
        )
        self.quality_banner.setStyleSheet(
            f"QLabel#QualityBanner {{"
            f"  background: {bg}; border: 1px solid {border};"
            f"  border-radius: 5px; color: #f0d080;"
            f"  font-size: 11px; font-weight: 600;"
            f"  padding: 6px 10px;"
            f"}}"
        )
        self.quality_banner.setVisible(True)

    def _update_ai_provider_hint(self, provider: str):
        info = AI_PROVIDERS.get(provider, {})
        self.ai_provider_hint.setText(info.get("desc", ""))

    def _run_ai_improve(self):
        text = self.editor.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Sin contenido", "El editor está vacío.")
            return
        provider = self.ai_provider_combo.currentText()
        hint = AI_PROVIDERS[provider]["key_hint"]
        supports_vision = AI_PROVIDERS[provider].get("vision", False)

        if provider not in self._api_keys or not self._api_keys[provider]:
            key, ok = QInputDialog.getText(
                self, f"API Key — {provider}",
                f"Ingresa tu API Key de {provider}\n(formato: {hint}):",
                QLineEdit.EchoMode.Password
            )
            if not ok or not key.strip():
                return
            self._api_keys[provider] = key.strip()

        quality = ocr_quality_score(text)
        if quality["level"] == "bad":
            vision_note = (
                "\n\nℹ️ Se enviará la imagen original a la IA para recuperar\n"
                "el texto que Tesseract no pudo leer." if supports_vision else ""
            )
            reply = QMessageBox.question(
                self, "Calidad muy baja",
                f"El texto tiene calidad OCR muy baja.{vision_note}\n\n¿Continuar?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply != QMessageBox.StandardButton.Yes:
                return

        self.btn_ai.setEnabled(False)
        vision_icon = "🖼️ " if supports_vision and self._current_image_path else ""
        self.btn_ai.setText(f"⏳  {vision_icon}Procesando IA…")
        self.lbl_status.setText(f"⏳ {provider} analizando imagen + texto…" if supports_vision
                                 else f"⏳ {provider} mejorando texto…")

        # v4.4: pasar imagen al hilo IA
        self._ai_thread = AIImproveThread(
            text, self._api_keys[provider], provider,
            image_path=self._current_image_path,
            parent=self
        )
        self._ai_thread.result_ready.connect(self._on_ai_result)
        self._ai_thread.error_occurred.connect(self._on_ai_error)
        self._ai_thread.start()

    def _on_ai_result(self, improved: str):
        old_text = self.editor.toPlainText().strip()
        if old_text:
            self._add_to_history("(pre-IA) " + self._last_title(), old_text)
        if len(improved) > MAX_CHARS:
            improved = improved[:MAX_CHARS]
        self.editor.setPlainText(improved)
        self.quality_banner.setVisible(False)
        self.btn_ai.setEnabled(True)
        self.btn_ai.setText("✨  Mejorar con IA")
        provider = self.ai_provider_combo.currentText()
        self.lbl_status.setText(f"✔ Texto reconstruido por {provider}")
        quality = ocr_quality_score(improved)
        self._show_quality_banner(quality)
        self._add_to_history("(IA) " + self._last_title(), improved)

    def _on_ai_error(self, error_msg: str):
        provider = self.ai_provider_combo.currentText()
        self.btn_ai.setEnabled(True)
        self.btn_ai.setText("✨  Mejorar con IA")
        self.lbl_status.setText("✗ Error IA")
        if "401" in error_msg or "authentication" in error_msg.lower() or "invalid" in error_msg.lower():
            self._api_keys.pop(provider, None)
            QMessageBox.critical(self, f"API Key inválida — {provider}",
                                 "La API Key fue rechazada.\n"
                                 "Se borrará para que puedas ingresar una nueva.")
        else:
            QMessageBox.critical(self, f"Error al mejorar con {provider}",
                                 f"No se pudo conectar con la API:\n{error_msg}")

    def _last_title(self) -> str:
        text = self.editor.toPlainText()
        m = re.search(r'^title:\s*(.+)$', text, re.MULTILINE)
        return m.group(1).strip() if m else "Nota"

    def _add_to_history(self, title: str, content: str):
        db_save_entry(self._db, title, content)
        self._refresh_history_widget()

    def _refresh_history_widget(self):
        self.history_list.clear()
        entries = db_load_all(self._db)
        for entry in entries:
            ts_short = entry["timestamp"][11:19]
            item = QListWidgetItem(f"{ts_short}  —  {entry['title']}")
            item.setData(Qt.ItemDataRole.UserRole, entry["content"])
            item.setToolTip(entry["timestamp"])
            self.history_list.addItem(item)

    def _load_from_history(self, item: QListWidgetItem):
        content = item.data(Qt.ItemDataRole.UserRole)
        if content:
            self.editor.setPlainText(content)
            for btn in (self.btn_md, self.btn_txt, self.btn_docx, self.btn_pdf):
                btn.setEnabled(True)
            self.btn_ai.setEnabled(True)

    def _clear_history(self):
        if QMessageBox.question(
            self, "Limpiar historial",
            "¿Eliminar todo el historial guardado?\nEsta acción no se puede deshacer.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        ) == QMessageBox.StandardButton.Yes:
            db_clear(self._db)
            self.history_list.clear()

    def _export(self, fmt: str):
        text = self.editor.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Sin contenido", "El editor está vacío.")
            return
        fecha = datetime.now().strftime("%Y-%m-%d")
        filters = {
            "md":   ("Markdown (*.md)",       f"{fecha}-nota.md"),
            "txt":  ("Texto plano (*.txt)",    f"{fecha}-nota.txt"),
            "docx": ("Word Document (*.docx)", f"{fecha}-nota.docx"),
            "pdf":  ("PDF (*.pdf)",            f"{fecha}-nota.pdf"),
        }
        filt, default_name = filters[fmt]
        path, _ = QFileDialog.getSaveFileName(self, f"Exportar {fmt.upper()}", default_name, filt)
        if not path:
            return
        if not path.lower().endswith(f".{fmt}"):
            path += f".{fmt}"
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        self.lbl_status.setText(f"⏳ Exportando {fmt.upper()}…")
        QApplication.processEvents()
        try:
            if fmt in ("md", "txt"):
                ok, err = export_to_txt(text, path)
            elif fmt == "docx":
                ok, err = export_to_docx(text, path)
            elif fmt == "pdf":
                ok, err = export_to_pdf(text, path)
            else:
                ok, err = False, "Formato desconocido"
        finally:
            QApplication.restoreOverrideCursor()
        if ok:
            self.lbl_status.setText(f"✔ {os.path.basename(path)}")
            QMessageBox.information(self, "Exportado", f"Archivo guardado:\n{path}")
        else:
            self.lbl_status.setText("✗ Error al exportar")
            QMessageBox.critical(self, f"Error al exportar {fmt.upper()}",
                                 f"No se pudo guardar el archivo:\n{err}")

    def _clear_editor(self):
        if self.editor.toPlainText().strip():
            if QMessageBox.question(
                self, "Limpiar editor", "¿Descartar el contenido actual?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            ) != QMessageBox.StandardButton.Yes:
                return
        self.editor.clear()
        self.quality_banner.setVisible(False)
        for btn in (self.btn_md, self.btn_txt, self.btn_docx, self.btn_pdf):
            btn.setEnabled(False)
        self.btn_ai.setEnabled(False)
        self.lbl_status.setText("en espera")

    def closeEvent(self, event):
        if hasattr(self, '_db') and self._db:
            self._db.close()
        super().closeEvent(event)

# ══════════════════════════════════════════════════════════════════
#  LANZADOR PROFESIONAL (CLI + GUI)
# ══════════════════════════════════════════════════════════════════

def mostrar_banner():
    """Imprime el banner estilo Livewire en la consola sin warnings"""
    # Usamos r"" (raw strings) para evitar el SyntaxWarning de las diagonales
    print("\033[94m" + r"  ____             _                _   _ _     _             " + "\033[0m")
    print("\033[94m" + r" / ___| _   _ _ __ | |_ __ ___  __  | | | (_)___(_) ___  _ __  " + "\033[0m")
    print("\033[96m" + r" \___ \| | | | '_ \| __/ _` \ \/ /  | | | | / __| |/ _ \| '_ \ " + "\033[0m")
    print("\033[96m" + r"  ___) | |_| | | | | || (_| |>  <   \ \_/ / \__ \ | (_) | | | |" + "\033[0m")
    print("\033[90m" + r" |____/ \__, |_| |_|\__\__,_/_/\_\   \___/|_|___/_|\___/|_| |_|" + "\033[0m")
    print("\033[90m" + r"        |___/                                                  " + "\033[0m")
    print("\n\033[32m  DONE \033[0m Component Started: \033[1mMainView\033[0m")
    print(f"\033[32m  INFO \033[0m Database Connected: \033[1m{DB_PATH}\033[0m\n")

def main_launcher():
    """Punto de entrada principal para el comando 'syntaxvision'"""
    mostrar_banner()
    
    # 1. Configurar Tesseract (Prioriza la carpeta del proyecto para portabilidad)
    exe_tesseract = resource_path(os.path.join("Tesseract-OCR", "tesseract.exe"))
    if os.path.isfile(exe_tesseract):
        pytesseract.pytesseract.tesseract_cmd = exe_tesseract
    else:
        # Fallback a instalación global si no está en la carpeta del clon
        for _candidate in [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]:
            if os.path.isfile(_candidate):
                pytesseract.pytesseract.tesseract_cmd = _candidate
                break

    # 2. Configurar ruta de modelos para EasyOCR relativa a la instalación
    os.environ["EASYOCR_MODULE_PATH"] = resource_path(os.path.join("resources", "models"))

    # 3. Iniciar la App
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    
    # Cargar icono global
    if os.path.isfile(LOGO_PATH):
        app.setWindowIcon(QIcon(LOGO_PATH))
    
    register_montserrat()
    
    window = SyntaxVision()
    window.show()
    sys.exit(app.exec())

# ══════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    main_launcher()

    